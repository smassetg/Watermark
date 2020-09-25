# This script automatically creates certificates from the names provided in the excel file
# this script works on all length names, and updates the certificates accordingly

'''
24/09/2020 Issue
os.remove doesn't work for all of the word docs and .pdf files from the working folder - unsure why
'''

import docx
from docx import Document
import PyPDF2
import docx2pdf
from docx2pdf import convert
import openpyxl as op
from docx.shared import Pt
import os
import shutil
import time

# Names list - this takes names in an excel file to insert on the certificates
names_list = 'C:/Users/sebas/OneDrive/Software/Python/Watermark/Additional_Names/Paid_Members_Certificate_List.xlsx'
# Worksheet variables declaration
wb1 = op.load_workbook(names_list)
ws1 = wb1.active
mr = ws1.max_row
mc = ws1.max_column

for i in range (1, mr+1):
    for j in range (1, mc+1):
        #This cycles through the names in the excel sheet, and saves in the 'new_name' variable
        new_name = ws1.cell(row = i, column = j).value
        print(new_name)
        name_length = len(new_name)

        #This opens up the watermark word document based on the length of the new_name string
        #This if statement now works and will select the correct certificate 'name' template
        if name_length <= 29:
            print(name_length, ' - name_length is under or equal to 29')
            doc = docx.Document('C:/Users/sebas/OneDrive/Software/Python/Watermark/New_Word_Docs/Supporting_Docs/Watermark_Development.docx')
            size = Pt(35)
        elif name_length > 32:
            print(name_length, ' - name_length is over to 32')
            doc = docx.Document('C:/Users/sebas/OneDrive/Software/Python/Watermark/New_Word_Docs/Supporting_Docs/Watermark_Development_over_32_chars_2.docx')
            size = Pt(30)
        else:
            print(name_length, ' - name_length is between 30 & 32')
            doc = docx.Document('C:/Users/sebas/OneDrive/Software/Python/Watermark/New_Word_Docs/Supporting_Docs/Watermark_Development_over_30_chars_2.docx')
            size = Pt(35)

    # This script needed to be indented out of the above if statement
    # it was not pulling through the correct documents for each option
    # when indented up a level it referenced the correct documents

    style = doc.styles['Normal']
    font = style.font
    #from docx.shared import Pt
    font.name = 'Calibri'
    #font.size = Pt(36)
    font.size = size
    font.bold = True

    #This iterates through the paragraphs in the word doc and updates the name with 'new_name'
    for paragraph in doc.paragraphs:
        if 'name' in paragraph.text:
            print(paragraph.text)
            paragraph.text = new_name
            #This tests the new paragraph characters
            #text = doc.paragraphs[7].text
            text = paragraph
            print('text =', text)
            print('paragraph =', paragraph)

            save_name = (new_name +'.docx')

            doc.save(save_name)
            print('word document saved')
            convert(save_name)
            print('word document converted to pdf')

for (dirname, dirs, files) in os.walk('.'):   #What is the fullstop here? Perhaps is means in the current directoy?
    for filename in files:
        if filename.endswith('.pdf'):
            certificate = open('C:/Users/sebas/OneDrive/Software/Python/Watermark/Final_Certificate.pdf', 'rb')
            pdfReader = PyPDF2.PdfFileReader(certificate)
            watermark = PyPDF2.PdfFileReader(open(filename, 'rb'))
            certificatePage1 = pdfReader.getPage(0)
            certificatePage1.mergePage(watermark.getPage(0))
            pdfWriter = PyPDF2.PdfFileWriter()
            pdfWriter.addPage(certificatePage1)
            name_update = filename[:-4]
            result = open(name_update+'_Certificate.pdf', 'wb')
            pdfWriter.write(result)
            certificate.close()
            result.close()


# This for loop deletes the docx files
for (dirname, dirs, files) in os.walk('.'):   #What is the fullstop here? Perhaps is means in the current directoy?
    for filename in files:
        # This if statement removes the word docs from the folder
        if filename.endswith('.docx'):
            print(filename)
            os.remove(filename)
            print(filename, ' - has been removed')

# This for loop moves the certificate pdf files to the final certificates folder
for (dirname, dirs, files) in os.walk('.'):
    for filename in files:
        if filename.endswith('.pdf'):
            try:
                if 'Certificate' in filename:
                    # This next line joins the path of the directory and the filename ready for shutil move
                    src = (os.path.join('C:/Users/sebas/OneDrive/Software/Python/Watermark/Additional_Names/', filename))
                    # This next line moves the file from the source path to the destination folder
                    shutil.move(src, 'C:/Users/sebas/OneDrive/Software/Python/Watermark/Additional_Names/Final_Certificates')
                    # Python is giving me an exception that a file already exists, but it has already moved it
                    # I am not too sure why this is, try to work around this with the try & except statement
            except:
                pass

time.sleep(10)

# This for loop deletes the remaining pdf files
for (dirname, dirs, files) in os.walk('.'):
    for filename in files:
        if filename.endswith('.pdf'):
            os.remove(filename)
            print(filename, ' - has been removed')
